import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import re
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_docx_from_md(md_file_path, output_docx_path):
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    def process_code_block(lines):
        code_text = "".join(lines)
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "F4F5F7")
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)

    def process_table(lines):
        parsed_rows = []
        for l in lines:
            if re.match(r'^\s*\|?\s*:?-+:?\s*\|', l):
                continue # Header separator
            parts = [c.strip() for c in l.strip().strip('|').split('|')]
            if parts and any(parts):
                parsed_rows.append(parts)
                
        if not parsed_rows:
            return
            
        num_cols = max(len(r) for r in parsed_rows)
        table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for r_idx, row_data in enumerate(parsed_rows):
            is_header = (r_idx == 0)
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < num_cols:
                    cell = table.cell(r_idx, c_idx)
                    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                    
                    if is_header:
                        set_cell_background(cell, "1E293B")
                    elif r_idx % 2 == 1:
                        set_cell_background(cell, "F8FAFC")
                    else:
                        set_cell_background(cell, "FFFFFF")
                        
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    
                    # Remove markdown bold tags inside cell
                    clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                    run = p.add_run(clean_text)
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9.5)
                    
                    if is_header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    else:
                        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                        if cell_text.startswith('**') or '[PASS]' in cell_text or '[FAIL]' in cell_text:
                            run.font.bold = True
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_formatted_paragraph(text, style_type="normal", level=1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        
        if style_type == "h1":
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            clean_text = text.lstrip('#').strip()
            run = p.add_run(clean_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Slate 900
        elif style_type == "h2":
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            clean_text = text.lstrip('#').strip()
            run = p.add_run(clean_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Blue 900
        elif style_type == "h3":
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            clean_text = text.lstrip('#').strip()
            run = p.add_run(clean_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB) # Blue 600
        elif style_type == "list":
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25 * level)
            parse_inline_formatting(p, text)
        else: # Normal
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            parse_inline_formatting(p, text)

    def parse_inline_formatting(paragraph, text):
        # Handle bullet points prefix
        clean_text = text
        if clean_text.strip().startswith('- ') or clean_text.strip().startswith('* '):
            clean_text = '• ' + clean_text.strip()[2:]
        elif re.match(r'^\s*-\s*\[[ xX]\]', clean_text):
            checked = '[X]' if '[x]' in clean_text.lower() else '[ ]'
            clean_text = checked + ' ' + re.sub(r'^\s*-\s*\[[ xX]\]', '', clean_text).strip()

        # Tokenize for bold **text** and inline code `code`
        pattern = r'(\*\*.*?\*\*|`.*?`)'
        tokens = re.split(pattern, clean_text)
        
        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**'):
                run = paragraph.add_run(token[2:-2])
                run.font.name = 'Calibri'
                run.font.bold = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            elif token.startswith('`') and token.endswith('`'):
                run = paragraph.add_run(token[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E) # Teal
            else:
                run = paragraph.add_run(token)
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    for line in lines:
        raw_line = line.rstrip('\n\r')
        
        # Check code block toggle
        if raw_line.strip().startswith('```'):
            if in_code_block:
                process_code_block(code_lines)
                code_lines = []
                in_code_block = False
            else:
                if in_table:
                    process_table(table_lines)
                    table_lines = []
                    in_table = False
                in_code_block = True
            continue
            
        if in_code_block:
            code_lines.append(raw_line + '\n')
            continue
            
        # Check table
        if raw_line.strip().startswith('|'):
            if not in_table:
                in_table = True
            table_lines.append(raw_line)
            continue
        elif in_table:
            process_table(table_lines)
            table_lines = []
            in_table = False
            
        if not raw_line.strip():
            continue
            
        if raw_line.startswith('# '):
            add_formatted_paragraph(raw_line, "h1")
        elif raw_line.startswith('## '):
            add_formatted_paragraph(raw_line, "h2")
        elif raw_line.startswith('### '):
            add_formatted_paragraph(raw_line, "h3")
        elif raw_line.strip().startswith('- ') or raw_line.strip().startswith('* ') or re.match(r'^\s*-\s*\[[ xX]\]', raw_line):
            add_formatted_paragraph(raw_line, "list")
        elif raw_line.strip() == '---':
            continue
        else:
            add_formatted_paragraph(raw_line, "normal")
            
    # Flush remaining blocks
    if in_code_block and code_lines:
        process_code_block(code_lines)
    if in_table and table_lines:
        process_table(table_lines)
        
    doc.save(output_docx_path)
    print(f"Successfully converted {md_file_path} to {output_docx_path}")

if __name__ == '__main__':
    md_path = os.path.join(os.path.dirname(__file__), '..', 'PRE_DEPLOYMENT_TESTING_DOC.md')
    docx_path = os.path.join(os.path.dirname(__file__), '..', 'PRE_DEPLOYMENT_TESTING_DOC.docx')
    create_docx_from_md(md_path, docx_path)
